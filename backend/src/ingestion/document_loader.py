import csv
import os
import time
from pathlib import Path
from typing import Any, List, Optional, Union
from dataclasses import dataclass, field
from loguru import logger


@dataclass
class Document:
    id: str
    title: str
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class LoadResult:
    success: bool
    document: Optional[Document] = None
    error: Optional[str] = None
    skipped: bool = False


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".txt": "text",
        ".md": "markdown",
        ".csv": "csv",
        ".pptx": "pptx",
        ".xlsx": "xlsx",
        ".xls": "xlsx",
        ".ppt": "pptx",
        ".html": "html",
        ".htm": "html",
        ".json": "json",
        ".xml": "xml",
        ".zip": "zip",
        ".epub": "epub",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".gif": "image",
        ".bmp": "image",
        ".tiff": "image",
        ".mp3": "audio",
        ".wav": "audio",
        ".m4a": "audio",
    }
    
    MARKITDOWN_EXTENSIONS = {
        ".pdf", ".docx", ".pptx", ".xlsx", ".xls", ".ppt",
        ".html", ".htm", ".json", ".xml", ".zip", ".epub",
        ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff",
        ".mp3", ".wav", ".m4a", ".csv", ".md", ".txt"
    }

    def __init__(
        self,
        max_retries: int = 3,
        retry_delay: float = 0.5,
        skip_empty: bool = True,
        min_content_length: int = 10,
        use_markitdown: bool = True,
    ):
        self.max_retries = max_retries
        self.retry_delay = retry_delay
        self.skip_empty = skip_empty
        self.min_content_length = min_content_length
        self.use_markitdown = use_markitdown
        self._markitdown_instance = None
        self.loaders = {
            "pdf": self._load_pdf,
            "docx": self._load_docx,
            "text": self._load_text,
            "markdown": self._load_text,
            "csv": self._load_csv,
            "pptx": self._load_with_markitdown,
            "xlsx": self._load_with_markitdown,
            "html": self._load_with_markitdown,
            "json": self._load_with_markitdown,
            "xml": self._load_with_markitdown,
            "zip": self._load_with_markitdown,
            "epub": self._load_with_markitdown,
            "image": self._load_with_markitdown,
            "audio": self._load_with_markitdown,
        }
    
    def _get_markitdown(self):
        if self._markitdown_instance is None:
            try:
                from markitdown import MarkItDown
                self._markitdown_instance = MarkItDown(enable_plugins=False)
            except ImportError:
                raise ImportError(
                    "markitdown is required for this file type. "
                    "Install with: pip install markitdown"
                )
        return self._markitdown_instance

    def _detect_encoding(self, file_path: Path) -> str:
        """Detect file encoding using chardet if available.

        Args:
            file_path: Path to the file

        Returns:
            Detected encoding or 'utf-8' as fallback
        """
        try:
            import chardet

            with open(file_path, "rb") as f:
                raw_data = f.read(10000)
                result = chardet.detect(raw_data)
                encoding = result.get("encoding")
                confidence = result.get("confidence", 0)
                if encoding and confidence > 0.7:
                    logger.debug(f"Detected encoding: {encoding} (confidence: {confidence})")
                    return encoding
        except ImportError:
            pass
        return "utf-8"

    def _retry_load(self, loader_func, file_path: Path, *args, **kwargs) -> Any:
        """Retry loading a file with exponential backoff.

        Args:
            loader_func: Loader function to call
            file_path: Path to the file
            *args: Additional positional arguments
            **kwargs: Additional keyword arguments

        Returns:
            Result of the loader function

        Raises:
            Exception: If all retry attempts fail
        """
        last_error = None
        for attempt in range(self.max_retries):
            try:
                return loader_func(file_path, *args, **kwargs)
            except Exception as e:
                last_error = e
                if attempt < self.max_retries - 1:
                    logger.warning(f"Load failed (attempt {attempt + 1}/{self.max_retries}): {e}")
                    time.sleep(self.retry_delay * (attempt + 1))
        raise last_error

    def load(self, file_path: Union[str, Path]) -> Document:
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        if self.use_markitdown and ext in self.MARKITDOWN_EXTENSIONS:
            return self._retry_load(self._load_with_markitdown, file_path)

        doc_type = self.SUPPORTED_EXTENSIONS[ext]
        loader = self.loaders.get(doc_type)

        if loader is None:
            raise ValueError(f"No loader for type: {doc_type}")

        return self._retry_load(loader, file_path)

    def load_safe(self, file_path: Union[str, Path]) -> LoadResult:
        """Load a document safely, returning a LoadResult instead of raising exceptions.

        Args:
            file_path: Path to the file

        Returns:
            LoadResult with success status and loaded document or error
        """
        try:
            doc = self.load(file_path)
            content = doc.content.strip() if doc.content else ""
            if self.skip_empty and len(content) < self.min_content_length:
                return LoadResult(
                    success=False,
                    skipped=True,
                    error=f"Content too short ({len(content)} chars), skipped",
                )
            return LoadResult(success=True, document=doc)
        except Exception as e:
            return LoadResult(success=False, error=str(e))

    def load_batch(self, directory: Union[str, Path]) -> List[Document]:
        """Load multiple documents from a directory.

        Args:
            directory: Directory to load documents from

        Returns:
            List of loaded documents
        """
        directory = Path(directory)
        documents = []
        errors = []

        for file_path in directory.rglob("*"):
            if file_path.is_file():
                ext = file_path.suffix.lower()
                if ext in self.SUPPORTED_EXTENSIONS:
                    result = self.load_safe(file_path)
                    if result.success and result.document:
                        documents.append(result.document)
                        logger.info(f"Loaded: {file_path.name}")
                    elif result.error and not result.skipped:
                        errors.append(f"{file_path.name}: {result.error}")
                        logger.error(f"Failed to load {file_path.name}: {result.error}")

        logger.info(f"Loaded {len(documents)} documents from {directory}")
        if errors:
            logger.warning(f"Failed to load {len(errors)} files: {errors}")
        return documents

    def load_batch_with_results(
        self, directory: Union[str, Path]
    ) -> tuple[List[Document], List[LoadResult]]:
        """Load multiple documents from a directory, returning both documents and results.

        Args:
            directory: Directory to load documents from

        Returns:
            Tuple of (loaded documents, load results)
        """
        directory = Path(directory)
        documents = []
        results = []

        for file_path in directory.rglob("*"):
            if file_path.is_file():
                ext = file_path.suffix.lower()
                if ext in self.SUPPORTED_EXTENSIONS:
                    result = self.load_safe(file_path)
                    results.append(result)
                    if result.success and result.document:
                        documents.append(result.document)
                        logger.info(f"Loaded: {file_path.name}")
                    else:
                        logger.error(f"Failed to load {file_path.name}: {result.error}")

        return documents, results

    def _load_pdf(self, file_path: Path) -> Document:
        """Load a PDF document.

        Args:
            file_path: Path to the PDF file

        Returns:
            Loaded document

        Raises:
            ImportError: If PyPDF2 is not installed
            ValueError: If no text content is extracted from PDF
        """
        try:
            import PyPDF2

            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)

                if reader.is_encrypted:
                    logger.warning(f"PDF is encrypted, attempting to decrypt: {file_path}")
                    try:
                        reader.decrypt("")
                    except Exception:
                        pass

                pages = []
                for page_num, page in enumerate(reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            pages.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        continue

                if not pages:
                    raise ValueError(f"No text content extracted from PDF: {file_path}")

                content = "\n\n".join(pages)
                metadata = reader.metadata or {}
                return Document(
                    id=self._generate_id(file_path),
                    title=file_path.stem,
                    content=content,
                    metadata={
                        "source": str(file_path),
                        "num_pages": len(reader.pages),
                        "num_text_pages": len(pages),
                        "pdf_metadata": {k: str(v) for k, v in metadata.items() if v},
                        "is_encrypted": reader.is_encrypted,
                    },
                )
        except ImportError:
            raise ImportError(
                "PyPDF2 is required for PDF loading. Install with: pip install PyPDF2"
            )
        except Exception as e:
            if "PDF has been corrupted" in str(e) or "malformed" in str(e).lower():
                raise ValueError(f"Corrupted PDF file: {file_path}. Error: {e}")
            raise

    def _load_docx(self, file_path: Path) -> Document:
        """加载 DOCX 文档

        Args:
            file_path: DOCX 文件路径

        Returns:
            加载的文档

        Raises:
            ImportError: 如果 python-docx 未安装
            ValueError: 如果 DOCX 文件中没有文本内容
        """
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

            content = "\n\n".join(paragraphs)
            if not content.strip():
                raise ValueError("No text content in DOCX file")

            return Document(
                id=self._generate_id(file_path),
                title=file_path.stem,
                content=content,
                metadata={
                    "source": str(file_path),
                    "num_paragraphs": len(paragraphs),
                    "num_tables": len(tables),
                    "has_tables": len(tables) > 0,
                },
            )
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX loading. Install with: pip install python-docx"
            )

    def _load_text(self, file_path: Path) -> Document:
        """加载文本文件（包括 TXT 和 Markdown）

        Args:
            file_path: 文本文件路径

        Returns:
            加载的文档

        Raises:
            ValueError: 如果文件为空
        """
        encoding = self._detect_encoding(file_path)
        try:
            with open(file_path, "r", encoding=encoding, errors="replace") as f:
                content = f.read()
        except UnicodeDecodeError:
            logger.warning(f"Encoding {encoding} failed, trying utf-8 with ignore")
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

        if not content.strip():
            raise ValueError(f"Empty file: {file_path}")

        file_type = "markdown" if file_path.suffix.lower() == ".md" else "text"
        return Document(
            id=self._generate_id(file_path),
            title=file_path.stem,
            content=content,
            metadata={
                "source": str(file_path),
                "file_type": file_type,
                "encoding": encoding,
            },
        )

    def _load_csv(self, file_path: Path) -> Document:
        """加载 CSV 文件

        Args:
            file_path: CSV 文件路径

        Returns:
            加载的文档

        Raises:
            ValueError: 如果 CSV 解析失败或文件没有数据行
        """
        encoding = self._detect_encoding(file_path)
        rows = []
        headers = []

        try:
            with open(file_path, "r", encoding=encoding, errors="replace") as f:
                sample = f.read(4096)
                delimiter = self._detect_delimiter(sample)
        except Exception:
            delimiter = ","

        try:
            with open(file_path, "r", encoding=encoding, errors="replace") as f:
                reader = csv.reader(f, delimiter=delimiter)
                headers = next(reader, [])
                for row in reader:
                    if any(cell.strip() for cell in row):
                        rows.append(row)
        except Exception as e:
            raise ValueError(f"Failed to parse CSV: {e}")

        if not rows:
            raise ValueError("CSV file has no data rows")

        content = self._format_csv(headers, rows)
        return Document(
            id=self._generate_id(file_path),
            title=file_path.stem,
            content=content,
            metadata={
                "source": str(file_path),
                "file_type": "csv",
                "num_rows": len(rows),
                "num_columns": len(headers) if headers else 0,
                "columns": headers,
                "delimiter": delimiter,
                "encoding": encoding,
            },
        )

    def _detect_delimiter(self, sample: str) -> str:
        """检测 CSV 文件的分隔符

        Args:
            sample: 文件样本内容

        Returns:
            检测到的分隔符
        """
        delimiters = [",", ";", "\t", "|"]
        counts = {}
        for d in delimiters:
            counts[d] = sample.count(d)
        return max(counts, key=counts.get) if counts else ","

    def _format_csv(self, headers: List[str], rows: List[List[str]]) -> str:
        """格式化 CSV 数据为文本，保留表格结构

        Args:
            headers: CSV 表头
            rows: CSV 数据行

        Returns:
            格式化后的文本内容，保留Markdown表格结构
        """
        content_lines = []
        if not headers and not rows:
            return ""

        if headers:
            header_line = f"| {' | '.join(str(h) if h else '' for h in headers)} |"
            separator_line = f"|{' | '.join(['---'] * len(headers))}|"
            content_lines.append(header_line)
            content_lines.append(separator_line)

        for row in rows:
            padded_row = row + [''] * (len(headers) - len(row)) if headers else row
            row_line = f"| {' | '.join(str(cell) if cell else '' for cell in padded_row)} |"
            content_lines.append(row_line)

        return "\n".join(content_lines)

    def _load_with_markitdown(self, file_path: Path) -> Document:
        md = self._get_markitdown()
        try:
            result = md.convert(str(file_path))
            content = result.text_content
            
            if not content or not content.strip():
                raise ValueError(f"No content extracted from {file_path}")
            
            ext = file_path.suffix.lower()
            file_type = self.SUPPORTED_EXTENSIONS.get(ext, "unknown")
            
            return Document(
                id=self._generate_id(file_path),
                title=file_path.stem,
                content=content,
                metadata={
                    "source": str(file_path),
                    "file_type": file_type,
                    "loader": "markitdown",
                },
            )
        except Exception as e:
            raise ValueError(f"Failed to convert {file_path} with markitdown: {e}")

    def _generate_id(self, file_path: Path) -> str:
        """根据文件路径生成唯一ID

        Args:
            file_path: 文件路径

        Returns:
            生成的唯一ID
        """
        import hashlib

        return hashlib.md5(str(file_path).encode()).hexdigest()


def load_document(file_path: Union[str, Path]) -> Document:
    """加载单个文档

    Args:
        file_path: 文件路径

    Returns:
        加载的文档
    """
    loader = DocumentLoader()
    return loader.load(file_path)


def load_documents_from_directory(directory: Union[str, Path]) -> List[Document]:
    """从目录加载多个文档

    Args:
        directory: 目录路径

    Returns:
        加载的文档列表
    """
    loader = DocumentLoader()
    return loader.load_batch(directory)
